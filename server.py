import http.server
import socketserver
import webbrowser
import os
import json
import re
import urllib.parse
import zipfile
import xml.etree.ElementTree as ET
import io
import hashlib
import datetime
from difflib import SequenceMatcher
import pypdf
from email.parser import BytesParser

PORT = 8000

def parse_multipart_bytes(body, boundary):
    primary_files = []
    secondary_files = []
    
    boundary_bytes = b"--" + boundary.encode('utf-8')
    chunks = body.split(boundary_bytes)
    
    for chunk in chunks:
        if not chunk or chunk.strip() == b"" or chunk.strip() == b"--":
            continue
            
        if chunk.startswith(b"\r\n"):
            chunk = chunk[2:]
        elif chunk.startswith(b"\n"):
            chunk = chunk[1:]
            
        if b"\r\n\r\n" not in chunk:
            continue
            
        header_bytes, file_bytes = chunk.split(b"\r\n\r\n", 1)
        
        if file_bytes.endswith(b"\r\n"):
            file_bytes = file_bytes[:-2]
        elif file_bytes.endswith(b"\n"):
            file_bytes = file_bytes[:-1]
            
        headers_lines = header_bytes.split(b"\r\n")
        content_disposition = b""
        for line in headers_lines:
            if line.lower().startswith(b"content-disposition:"):
                content_disposition = line
                break
                
        if not content_disposition:
            continue
            
        name_match = re.search(b'name="([^"]+)"', content_disposition, re.IGNORECASE)
        filename_match = re.search(b'filename="([^"]+)"', content_disposition, re.IGNORECASE)
        
        if name_match and filename_match:
            try:
                name_val = name_match.group(1).decode('utf-8')
            except Exception:
                name_val = name_match.group(1).decode('latin-1')
            try:
                filename_val = filename_match.group(1).decode('utf-8')
            except Exception:
                filename_val = filename_match.group(1).decode('latin-1')
            
            if name_val == 'primary_files':
                primary_files.append((filename_val, file_bytes))
            elif name_val == 'secondary_files':
                secondary_files.append((filename_val, file_bytes))
                
    return primary_files, secondary_files



PRESETS = [
    {
        "id": "sys_meta",
        "name": "文档元数据一致性审查",
        "description": "比对两份标书的PDF元数据（如作者、创建工具、修改时间等），检测是否使用相同排版环境编译。",
        "category": "系统预置",
        "enabled": True
    },
    {
        "id": "sys_image",
        "name": "图片资源二进制哈希审查",
        "description": "提取PDF中的图片指纹，判定两份标书中是否有完全一致的拓扑插图或架构图文件。",
        "category": "系统预置",
        "enabled": True
    },
    {
        "id": "sys_text",
        "name": "大段落文本相似度审查",
        "description": "基于 15-gram 中文分词滑动比对算法，检测两份标书在定制编写的非标书模板条款中的相似度。",
        "category": "系统预置",
        "enabled": True
    },
    {
        "id": "sys_template",
        "name": "标书模板遗留条款降权审查",
        "description": "智能匹配并过滤通用招标文件模板（如电网 boilerplate 官方用语），对合理雷同项进行自动降权降噪。",
        "category": "系统预置",
        "enabled": True
    },
    {
        "id": "sys_company",
        "name": "交叉公司简称泄露审查",
        "description": "交叉反查我方与对手方标书中的公司简称，判定是否存在“套模板时未删干净”的公司名铁证。",
        "category": "系统预置",
        "enabled": True
    }
]

DB_FILE = os.path.expanduser("~/.baoyu-skills/rules_db.json")

def load_rules_db():
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if "current_rules" in data:
                    db = {
                        "active_business_id": "biz_default",
                        "businesses": {
                            "biz_default": {
                                "id": "biz_default",
                                "name": "业务1",
                                "is_default": True,
                                "created_at": timestamp,
                                "updated_at": timestamp,
                                "rules": data["current_rules"]
                            }
                        }
                    }
                    save_rules_db(db)
                    return db
                return data
        except Exception:
            pass
    db = {
        "active_business_id": "biz_default",
        "businesses": {
            "biz_default": {
                "id": "biz_default",
                "name": "业务1",
                "is_default": True,
                "created_at": timestamp,
                "updated_at": timestamp,
                "rules": list(PRESETS)
            }
        }
    }
    save_rules_db(db)
    return db

def save_rules_db(db):
    try:
        os.makedirs(os.path.dirname(DB_FILE), exist_ok=True)
        with open(DB_FILE, 'w', encoding='utf-8') as f:
            json.dump(db, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

HISTORY_FILE = os.path.expanduser("~/.baoyu-skills/history_db.json")
LATEST_COLLUSION_DATA = {}


def load_history_db():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return []

def save_history_db(db):
    try:
        os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(db, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

def check_with_gemini(new_rule, current_rules, api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    prompt = f"""
    分析以下“新规则”与“现有规则列表”，判断新规则是否与现有规则存在“重复”(duplicate, 意思高度一致或极其相似) 或 “冲突”(conflict, 正反含义矛盾或对立)。
    
    现有规则列表：
    {json.dumps(current_rules, ensure_ascii=False, indent=2)}
    
    新规则：
    {json.dumps(new_rule, ensure_ascii=False, indent=2)}
    
    请只返回 JSON 数据，格式如下。如果无冲突和重复，请返回空数组。不要包含 markdown 代码块：
    [
      {{
        "incoming_id": "新规则ID",
        "existing_id": "现有规则ID",
        "existing_name": "现有规则名称",
        "type": "duplicate" 或者是 "conflict",
        "reason": "中文解释发生重复或冲突的具体原因，简短明了"
      }}
    ]
    """
    
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "responseMimeType": "application/json"
        }
    }
    
    req = urllib.request.Request(url, data=json.dumps(body).encode('utf-8'), headers=headers, method='POST')
    with urllib.request.urlopen(req, timeout=8) as response:
        res_data = json.loads(response.read().decode('utf-8'))
        text_response = res_data['candidates'][0]['content']['parts'][0]['text'].strip()
        if text_response.startswith("```json"):
            text_response = text_response[7:]
        if text_response.endswith("```"):
            text_response = text_response[:-3]
        return json.loads(text_response.strip())

def check_duplicate_and_conflict(new_rule, current_rules, api_key=None):
    if api_key:
        try:
            res = check_with_gemini(new_rule, current_rules, api_key)
            if res:
                return res
        except Exception as e:
            print(f"Gemini check failed: {e}")
            
    matches = []
    new_desc = new_rule.get('description', '').strip()
    
    for r in current_rules:
        r_desc = r.get('description', '').strip()
        
        sim = SequenceMatcher(None, new_desc, r_desc).ratio()
        if sim > 0.75:
            matches.append({
                "incoming_id": new_rule.get('id'),
                "existing_id": r.get('id'),
                "existing_name": r.get('name'),
                "type": "duplicate",
                "reason": f"规则内容与现有规则相似度高达 {sim*100:.1f}%"
            })
            continue
            
        topics = ["元数据", "图片", "文本", "相似度", "公司简称", "作者", "修改时间", "哈希", "模板", "格式"]
        matched_topic = None
        for t in topics:
            if t in new_desc and t in r_desc:
                matched_topic = t
                break
                
        if matched_topic:
            neg_words = ["禁止", "不能", "不许", "不可", "不得", "不要", "不一致", "不相同", "免除", "忽略"]
            new_neg = any(w in new_desc for w in neg_words)
            r_neg = any(w in r_desc for w in neg_words)
            if new_neg != r_neg:
                matches.append({
                    "incoming_id": new_rule.get('id'),
                    "existing_id": r.get('id'),
                    "existing_name": r.get('name'),
                    "type": "conflict",
                    "reason": f"针对【{matched_topic}】，此规则与现有规则 '{r.get('name')}' 的正负要求相左"
                })
    return matches

def extract_rules_from_file(filename, file_bytes):
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    if ext in ['.md', '.txt']:
        text = file_bytes.decode('utf-8', errors='ignore')
    elif ext == '.docx':
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx:
                xml_content = docx.read('word/document.xml')
                root = ET.fromstring(xml_content)
                texts = []
                for elem in root.iter():
                    if elem.tag.endswith('t'):
                        texts.append(elem.text)
                text = "".join(texts)
        except Exception as e:
            text = f"解析 docx 失败: {str(e)}"
    elif ext == '.pdf':
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            text = "".join(texts)
        except Exception as e:
            text = f"解析 pdf 失败: {str(e)}"
            
    parts = re.split(r'\n+|- |\* |\b\d+[\.\、\s]', text)
    rules = []
    for p in parts:
        cleaned = p.strip()
        if len(cleaned) > 8 and not cleaned.startswith('#'):
            rule_id = "rule_" + hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]
            name = cleaned[:16] + ("..." if len(cleaned) > 16 else "")
            rules.append({
                "id": rule_id,
                "name": name,
                "description": cleaned,
                "category": "导入规则",
                "enabled": True
            })
    return rules

def simplify_bidder_name(name):
    if "华路" in name:
        return "华路"
    if "至臻云" in name:
        return "至臻云"
    if "卓维" in name:
        return "卓维"
    if "数据易" in name:
        return "数据易"
    cleaned = re.sub(r'^(北京|广东|上海|深圳|天津|重庆|四川|江苏|浙江|山东|福建|湖北|湖南|河南|河北|安徽|江西|陕西|辽宁|吉林|黑龙江|山西|甘肃|青海|海南|云南|贵州|广西|内蒙古|西藏|宁夏|新疆)', '', name)
    cleaned = re.sub(r'(股份有限公司|有限责任公司|有限公司|信息技术|智能科技|网络)$', '', cleaned)
    return cleaned

def check_collusion_compare_wrapper(primary_path, secondary_path):
    from collections import defaultdict
    import sys
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    import check_collusion
    
    issues = []
    
    # 1. Compare metadata
    meta_matches = check_collusion.compare_metadata(primary_path, secondary_path)
    for m in meta_matches:
        try:
            pri_reader = pypdf.PdfReader(primary_path)
            sec_reader = pypdf.PdfReader(secondary_path)
            pri_meta = json.dumps(pri_reader.metadata or {}, ensure_ascii=False, indent=2)
            sec_meta = json.dumps(sec_reader.metadata or {}, ensure_ascii=False, indent=2)
        except:
            pri_meta, sec_meta = "", ""
        issues.append({
            "type": "META",
            "sec_page": m["secondary_pages"],
            "pri_page": m["primary_pages"],
            "rate": m["suspicion_rate"],
            "description": m["description"],
            "pri_text": f"主标书元数据:\n{pri_meta}",
            "sec_text": f"对照标书元数据:\n{sec_meta}"
        })
        
    # 2. Compare images
    image_matches = check_collusion.compare_images(primary_path, secondary_path)
    for m in image_matches:
        desc = m["description"].replace("[配图一致]", f"[配图一致] (主标{m['primary_pages']} - 陪标{m['secondary_pages']})")
        issues.append({
            "type": "IMAGE",
            "sec_page": m["secondary_pages"],
            "pri_page": m["primary_pages"],
            "rate": m["suspicion_rate"],
            "description": desc,
            "pri_text": f"图片资源哈希一致，出现在主标书的: {m['primary_pages']}",
            "sec_text": f"图片资源哈希一致，出现在对照标书的: {m['secondary_pages']}"
        })
        
    # 3. Compare text pages
    try:
        pri_reader = pypdf.PdfReader(primary_path)
        sec_reader = pypdf.PdfReader(secondary_path)
    except:
        return issues
        
    primary_pages_clean = []
    primary_pages_raw = []
    for page in pri_reader.pages:
        raw = page.extract_text() or ""
        primary_pages_raw.append(raw)
        primary_pages_clean.append(check_collusion.clean_text(raw))
        
    secondary_pages_clean = []
    secondary_pages_raw = []
    for page in sec_reader.pages:
        raw = page.extract_text() or ""
        secondary_pages_raw.append(raw)
        secondary_pages_clean.append(check_collusion.clean_text(raw))
        
    sec_ngram_index = defaultdict(list)
    ngram_size = 15
    for sec_idx, sec_text in enumerate(secondary_pages_clean):
        sec_page_num = sec_idx + 1
        if len(sec_text) < ngram_size:
            continue
        for i in range(len(sec_text) - ngram_size + 1):
            ngram = sec_text[i:i+ngram_size]
            sec_ngram_index[ngram].append(sec_page_num)
            
    candidates = defaultdict(int)
    for pri_idx, pri_text in enumerate(primary_pages_clean):
        pri_page_num = pri_idx + 1
        if len(pri_text) < ngram_size:
            continue
        for i in range(len(pri_text) - ngram_size + 1):
            ngram = pri_text[i:i+ngram_size]
            if ngram in sec_ngram_index:
                for sec_page_num in sec_ngram_index[ngram]:
                    candidates[(pri_page_num, sec_page_num)] += 1
                    
    min_match_len = 20
    for (pri_page, sec_page), common_count in candidates.items():
        if common_count < 3:
            continue
            
        pri_text_clean = primary_pages_clean[pri_page - 1]
        sec_text_clean = secondary_pages_clean[sec_page - 1]
        
        matcher = SequenceMatcher(None, pri_text_clean, sec_text_clean)
        matching_blocks = matcher.get_matching_blocks()
        
        total_match_len = 0
        longest_match = ""
        for block in matching_blocks:
            if block.size >= min_match_len:
                total_match_len += block.size
                match_str = pri_text_clean[block.a : block.a + block.size]
                if len(match_str) > len(longest_match):
                    longest_match = match_str
                    
        if total_match_len > 0:
            min_len = min(len(pri_text_clean), len(sec_text_clean))
            susp_rate = total_match_len / min_len if min_len > 0 else 0
            
            if susp_rate >= 0.10:
                match_type = check_collusion.classify_match(longest_match)
                
                if match_type == "标书模板内容未删除":
                    reported_type = "TEMPLATE"
                    desc = f"[模板未删] (主标P{pri_page} - 陪标P{sec_page}) 模板匹配率较高。最长匹配: '{longest_match[:25]}...'"
                    reported_susp_rate = susp_rate * 0.10
                else:
                    reported_type = "TEXT"
                    desc = f"[文本抄袭] (主标P{pri_page} - 陪标P{sec_page}) 文本重合率高。最长匹配: '{longest_match[:25]}...'"
                    reported_susp_rate = susp_rate
                    
                issues.append({
                    "type": reported_type,
                    "sec_page": f"P{sec_page}",
                    "pri_page": f"P{pri_page}",
                    "rate": reported_susp_rate,
                    "description": desc,
                    "pri_text": primary_pages_raw[pri_page - 1],
                    "sec_text": secondary_pages_raw[sec_page - 1]
                })
                
    issues = sorted(issues, key=lambda x: x["rate"], reverse=True)
    for idx, iss in enumerate(issues, 1):
        iss["idx"] = idx
        
    return issues

def generate_docx_report(competitor, data):

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    import io
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    title = doc.add_heading("投标自查诊断报告", level=1)
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        run.font.bold = True
        
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"报告生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").font.size = Pt(9.5)
    p_meta.add_run(f"对照排查单位: {competitor}").font.size = Pt(9.5)
    p_meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p_meta.runs[1].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    h_sum = doc.add_heading("1. 风险评估摘要", level=2)
    for run in h_sum.runs:
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        
    doc.add_paragraph(data.get("summary", ""))
    
    h_det = doc.add_heading("2. 详细排查点列表", level=2)
    for run in h_det.runs:
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        
    issues = data.get("issues", [])
    if not issues:
        doc.add_paragraph("未检测到明显的串标雷同项。")
    else:
        table_data = [["序号", "陪标单位", "陪标页码", "主标页码", "嫌疑率", "问题描述"]]
        for idx, iss in enumerate(issues, 1):
            rate_str = f"{iss.get('rate', 0) * 100:.1f}%"
            desc = re.sub(r'\*\*(.*?)\*\*', r'\1', iss.get('description', ''))
            table_data.append([
                str(idx),
                competitor,
                str(iss.get('sec_page', '')),
                str(iss.get('pri_page', '')),
                rate_str,
                desc
            ])
            
        rows = len(table_data)
        cols = 6
        table = doc.add_table(rows=rows, cols=cols)
        
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        
        col_widths = [Inches(0.5), Inches(1.0), Inches(0.8), Inches(0.8), Inches(1.0), Inches(2.4)]
        for r_idx, row_data in enumerate(table_data):
            row = table.rows[r_idx]
            for c_idx, cell_value in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.width = col_widths[c_idx]
                p = cell.paragraphs[0]
                run = p.add_run(cell_value)
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                
                if r_idx == 0:
                    run.font.bold = True
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    tcPr.append(shd)
                elif r_idx % 2 == 1:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FAFAFA"/>')
                    tcPr.append(shd)
                    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_html_report(competitor, data):
    import html
    rows_html = ""
    issues = data.get("issues", [])
    if not issues:
        rows_html = '<tr><td colspan="6" style="text-align: center; color: #6b7280;">未检测到明显的串标雷同项。</td></tr>'
    else:
        for idx, iss in enumerate(issues, 1):
            rate = iss.get('rate', 0)
            rate_str = f"{rate * 100:.1f}%"
            if rate >= 0.8:
                badge = f'<span class="badge badge-high">{rate_str}</span>'
            elif rate >= 0.3:
                badge = f'<span class="badge badge-medium">{rate_str}</span>'
            else:
                badge = f'<span class="badge badge-low">{rate_str}</span>'
                
            desc = html.escape(iss.get('description', ''))
            desc = desc.replace("**", "<b>").replace("**", "</b>")
            
            rows_html += f"""
            <tr>
                <td>{idx}</td>
                <td>{html.escape(competitor)}</td>
                <td>{html.escape(iss.get('sec_page', ''))}</td>
                <td>{html.escape(iss.get('pri_page', ''))}</td>
                <td>{badge}</td>
                <td>{desc}</td>
            </tr>
            """
            
    time_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    summary_escaped = html.escape(data.get("summary", "")).replace("\n", "<br/>")
    
    html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>标书自查诊断报告 - {html.escape(competitor)}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            color: #1f2937;
            background-color: #f9fafb;
            margin: 0;
            padding: 40px 20px;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background-color: white;
            padding: 40px;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            border: 1px solid #e5e7eb;
        }}
        h1 {{
            color: #111827;
            margin-top: 0;
            font-size: 28px;
            border-bottom: 2px solid #3b82f6;
            padding-bottom: 12px;
        }}
        .meta {{
            color: #6b7280;
            font-size: 14px;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #374151;
            font-size: 20px;
            margin-top: 30px;
            border-bottom: 1px solid #e5e7eb;
            padding-bottom: 8px;
        }}
        .summary-card {{
            background-color: #eff6ff;
            border-left: 4px solid #3b82f6;
            padding: 16px;
            border-radius: 4px;
            margin-bottom: 30px;
            font-size: 15px;
            line-height: 1.6;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
        }}
        th, td {{
            text-align: left;
            padding: 12px;
            border-bottom: 1px solid #e5e7eb;
            font-size: 14px;
        }}
        th {{
            background-color: #f3f4f6;
            font-weight: 600;
            color: #374151;
        }}
        tr:nth-child(even) td {{
            background-color: #fafafa;
        }}
        .badge {{
            display: inline-block;
            padding: 2px 8px;
            border-radius: 9999px;
            font-size: 12px;
            font-weight: 500;
        }}
        .badge-high {{ background-color: #fee2e2; color: #991b1b; }}
        .badge-medium {{ background-color: #fef3c7; color: #92400e; }}
        .badge-low {{ background-color: #d1fae5; color: #065f46; }}
        
        @media print {{
            body {{ background-color: white; padding: 0; }}
            .container {{ box-shadow: none; border: none; padding: 0; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>投标自查诊断报告</h1>
        <div class="meta">
            <div><strong>生成时间:</strong> {time_str}</div>
            <div><strong>对照单位:</strong> {html.escape(competitor)}</div>
        </div>
        
        <h2>1. 风险评估摘要</h2>
        <div class="summary-card">
            {summary_escaped}
        </div>
        
        <h2>2. 详细排查点列表</h2>
        <table>
            <thead>
                <tr>
                    <th style="width: 50px;">序号</th>
                    <th>陪标单位</th>
                    <th style="width: 80px;">陪标页码</th>
                    <th style="width: 80px;">主标页码</th>
                    <th style="width: 80px;">嫌疑率</th>
                    <th>问题描述</th>
                </tr>
            </thead>
            <tbody>
                {rows_html}
            </tbody>
        </table>
    </div>
</body>
</html>"""
    return html_template

class CustomHandler(http.server.SimpleHTTPRequestHandler):
    def do_GET(self):
        if self.path.startswith('/api/get_key'):
            query = urllib.parse.urlparse(self.path).query
            params = urllib.parse.parse_qs(query)
            env_key = params.get('env_key', [''])[0]
            
            val = self.load_key(env_key)
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'key': val}).encode('utf-8'))
            return
        elif self.path == '/api/get_rules':
            db = load_rules_db()
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps(db).encode('utf-8'))
            return
        elif self.path == '/api/get_history':
            history = load_history_db()
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps(history).encode('utf-8'))
            return
        elif self.path.startswith('/api/export_report'):
            query = urllib.parse.urlparse(self.path).query
            params = urllib.parse.parse_qs(query)
            fmt = params.get('format', ['word'])[0].lower()
            competitor = params.get('competitor', [''])[0]
            
            global LATEST_COLLUSION_DATA
            
            collusion_data = None
            if competitor and competitor in LATEST_COLLUSION_DATA:
                collusion_data = LATEST_COLLUSION_DATA
            elif LATEST_COLLUSION_DATA:
                collusion_data = LATEST_COLLUSION_DATA
            elif competitor:
                history = load_history_db()
                for record in history:
                    if record.get('collusionData') and competitor in record['collusionData']:
                        collusion_data = record['collusionData']
                        break
            
            if not collusion_data:
                history = load_history_db()
                if history and history[0].get('collusionData'):
                    collusion_data = history[0]['collusionData']
            
            if not collusion_data:
                self.send_response(400)
                self.send_header('Content-type', 'text/html; charset=utf-8')
                self.end_headers()
                self.wfile.write("错误: 未找到排查数据，请重新排查后重试。".encode('utf-8'))
                return
                
            import io
            import zipfile
            
            zip_buffer = io.BytesIO()
            try:
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for comp_name, comp_data in collusion_data.items():
                        if fmt == 'word':
                            file_bytes = generate_docx_report(comp_name, comp_data)
                            file_name = f"诊断报告_{comp_name}.docx"
                        elif fmt == 'md':
                            lines = [
                                f"# 投标自查诊断报告 - {comp_name}",
                                "",
                                f"- **报告生成时间**: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                                f"- **对照排查单位**: {comp_name}",
                                "",
                                "## 1. 风险评估摘要",
                                "",
                                comp_data.get("summary", ""),
                                "",
                                "## 2. 详细排查点列表",
                                "",
                                "| 序号 | 陪标单位 | 陪标页码 | 主标页码 | 嫌疑疑似率 | 问题描述 |",
                                "| --- | --- | --- | --- | --- | --- |"
                            ]
                            for idx, iss in enumerate(comp_data.get("issues", []), 1):
                                rate_str = f"{iss.get('rate', 0) * 100:.1f}%"
                                desc = iss.get('description', '')
                                lines.append(f"| {idx} | {comp_name} | {iss.get('sec_page', '')} | {iss.get('pri_page', '')} | {rate_str} | {desc} |")
                            file_bytes = "\n".join(lines).encode('utf-8')
                            file_name = f"诊断报告_{comp_name}.md"
                        elif fmt == 'html' or fmt == 'pdf':
                            html_text = generate_html_report(comp_name, comp_data)
                            if fmt == 'pdf':
                                print_script = "<script>window.onload = function() { window.print(); }</script></body>"
                                html_text = html_text.replace("</body>", print_script)
                            file_bytes = html_text.encode('utf-8')
                            file_name = f"诊断报告_{comp_name}.html"
                        elif fmt == 'csv':
                            lines = ["\ufeff序号,陪标单位,陪标页码,主标页码,嫌疑率,问题描述"]
                            for idx, iss in enumerate(comp_data.get("issues", []), 1):
                                rate_str = f"{iss.get('rate', 0) * 100:.1f}%"
                                desc_escaped = iss.get('description', '').replace('"', '""')
                                lines.append(f'{idx},{comp_name},{iss.get("sec_page", "")},{iss.get("pri_page", "")},{rate_str},"{desc_escaped}"')
                            file_bytes = "\n".join(lines).encode('utf-8')
                            file_name = f"诊断报告_{comp_name}.csv"
                        else:
                            continue
                        zip_file.writestr(file_name, file_bytes)
                        
                zip_bytes = zip_buffer.getvalue()
                
                self.send_response(200)
                filename_escaped = urllib.parse.quote("自查排查诊断报告.zip")
                self.send_header('Content-type', 'application/zip')
                self.send_header('Content-Disposition', f"attachment; filename*=UTF-8''{filename_escaped}")
                self.end_headers()
                self.wfile.write(zip_bytes)
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'text/html; charset=utf-8')
                self.end_headers()
                self.wfile.write(f"打包导出报告失败: {e}".encode('utf-8'))
            return
        return super().do_GET()

    def do_POST(self):
        if self.path == '/api/save_key':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            env_key = data.get('env_key', '')
            new_val = data.get('key', '')
            
            success, err = self.save_key(env_key, new_val)
            
            self.send_response(200 if success else 500)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'success': success, 'error': err}).encode('utf-8'))
            return
        elif self.path == '/api/save_rules':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            db = load_rules_db()
            active_id = db.get("active_business_id", "biz_default")
            
            if active_id in db["businesses"]:
                db["businesses"][active_id]["rules"] = data.get("current_rules", [])
                db["businesses"][active_id]["updated_at"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/create_business':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            name = data.get("name", "未命名业务").strip()
            
            db = load_rules_db()
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            biz_id = "biz_" + datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            db["businesses"][biz_id] = {
                "id": biz_id,
                "name": name,
                "is_default": False,
                "created_at": timestamp,
                "updated_at": timestamp,
                "rules": list(PRESETS)
            }
            db["active_business_id"] = biz_id
            save_rules_db(db)
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True, "active_business_id": biz_id}).encode('utf-8'))
            return
        elif self.path == '/api/delete_business':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            biz_id = data.get("business_id")
            
            db = load_rules_db()
            is_default = db["businesses"].get(biz_id, {}).get("is_default", False)
            if len(db["businesses"]) <= 1 or is_default:
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": "默认集或最后一个业务集禁止删除"}).encode('utf-8'))
                return
                
            if biz_id in db["businesses"]:
                del db["businesses"][biz_id]
                if db["active_business_id"] == biz_id:
                    db["active_business_id"] = list(db["businesses"].keys())[0]
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True, "active_business_id": db["active_business_id"]}).encode('utf-8'))
            return
        elif self.path == '/api/rename_business':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            biz_id = data.get("business_id")
            new_name = data.get("name", "未命名业务").strip()
            
            db = load_rules_db()
            if biz_id in db["businesses"] and new_name:
                db["businesses"][biz_id]["name"] = new_name
                db["businesses"][biz_id]["updated_at"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/switch_business':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            biz_id = data.get("business_id")
            
            db = load_rules_db()
            if biz_id in db["businesses"]:
                db["active_business_id"] = biz_id
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/set_default_business':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            biz_id = data.get("business_id")
            
            db = load_rules_db()
            if biz_id in db["businesses"]:
                for bid in db["businesses"]:
                    db["businesses"][bid]["is_default"] = (bid == biz_id)
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/reset_business_rules':
            db = load_rules_db()
            active_id = db.get("active_business_id", "biz_default")
            
            if active_id in db["businesses"]:
                db["businesses"][active_id]["rules"] = list(PRESETS)
                db["businesses"][active_id]["updated_at"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/import_rules':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            filename = data.get("filename", "")
            content_base64 = data.get("content", "")
            text_input = data.get("text", "")
            
            rules_to_check = []
            if content_base64:
                import base64
                file_bytes = base64.b64decode(content_base64)
                rules_to_check = extract_rules_from_file(filename, file_bytes)
            elif text_input:
                rule_id = "rule_" + hashlib.md5(text_input.encode('utf-8')).hexdigest()[:8]
                name = text_input[:16] + ("..." if len(text_input) > 16 else "")
                rules_to_check = [{
                    "id": rule_id,
                    "name": name,
                    "description": text_input,
                    "category": "导入规则",
                    "enabled": True
                }]
                
            db = load_rules_db()
            active_id = db.get("active_business_id", "biz_default")
            current_rules = db["businesses"].get(active_id, {}).get("rules", [])
            
            api_key = self.load_key("GEMINI_API_KEY")
            
            results = []
            for nr in rules_to_check:
                conflicts = check_duplicate_and_conflict(nr, current_rules, api_key)
                if conflicts:
                    nr["conflict_or_duplicate"] = conflicts[0]
                results.append(nr)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True, "rules": results}).encode('utf-8'))
            return
        elif self.path == '/api/resolve_imported_rules':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            resolutions = data.get("resolutions", [])
            
            db = load_rules_db()
            active_id = db.get("active_business_id", "biz_default")
            
            if active_id in db["businesses"]:
                current_rules = db["businesses"][active_id]["rules"]
                
                for item in resolutions:
                    rule = item.get("rule")
                    action = item.get("action")
                    
                    if not rule:
                        continue
                        
                    if action == "ignore":
                        continue
                    elif action == "overwrite":
                        conflict_info = rule.get("conflict_or_duplicate")
                        if conflict_info:
                            exist_id = conflict_info.get("existing_id")
                            current_rules = [r for r in current_rules if r.get("id") != exist_id]
                        if "conflict_or_duplicate" in rule:
                            del rule["conflict_or_duplicate"]
                        current_rules.append(rule)
                    else:
                        if "conflict_or_duplicate" in rule:
                            del rule["conflict_or_duplicate"]
                        current_rules.append(rule)
                        
                db["businesses"][active_id]["rules"] = current_rules
                db["businesses"][active_id]["updated_at"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                save_rules_db(db)
                
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/save_history':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            record = json.loads(post_data.decode('utf-8'))
            
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            record["timestamp"] = timestamp
            record["id"] = "hist_" + datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            db = load_history_db()
            db.insert(0, record)
            db = db[:50]
            save_history_db(db)
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"success": True}).encode('utf-8'))
            return
        elif self.path == '/api/analyze':
            content_type = self.headers.get('Content-Type')
            if not content_type or not content_type.startswith('multipart/form-data'):
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": "Invalid content type"}).encode('utf-8'))
                return
                
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            
            boundary = ""
            for param in content_type.split(';'):
                param = param.strip()
                if param.startswith('boundary='):
                    boundary = param.split('=', 1)[1]
                    if boundary.startswith('"') and boundary.endswith('"'):
                        boundary = boundary[1:-1]
            
            if not boundary:
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": "Missing boundary in Content-Type"}).encode('utf-8'))
                return
                
            primary_files, secondary_files = parse_multipart_bytes(body, boundary)
            
            if not primary_files or not secondary_files:
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": "请先上传我方标书及对照标书！"}).encode('utf-8'))
                return
                
            import tempfile
            import shutil
            
            temp_dir = tempfile.mkdtemp(dir=os.path.dirname(os.path.abspath(__file__)))
            
            try:
                primary_paths = []
                for fname, fbytes in primary_files:
                    path = os.path.join(temp_dir, "pri_" + fname)
                    with open(path, 'wb') as f:
                        f.write(fbytes)
                    primary_paths.append(path)
                    
                secondary_paths = []
                for fname, fbytes in secondary_files:
                    path = os.path.join(temp_dir, "sec_" + fname)
                    with open(path, 'wb') as f:
                        f.write(fbytes)
                    competitor_name = simplify_bidder_name(os.path.splitext(fname)[0])
                    secondary_paths.append((competitor_name, path))
                    
                collusion_data = {}
                pri_path = primary_paths[0]
                
                for competitor_name, sec_path in secondary_paths:
                    issues = check_collusion_compare_wrapper(pri_path, sec_path)
                    
                    max_rate = max([iss["rate"] for iss in issues]) if issues else 0
                    if max_rate >= 0.8:
                        summary = f"共检测到 {len(issues)} 个排查点，存在严重的定制文本大面积雷同及配图一致，具有极高的围标串标嫌疑。"
                    elif max_rate >= 0.3:
                        summary = f"共检测到 {len(issues)} 个排查点，存在部分雷同内容，具有中等围标串标嫌疑。"
                    else:
                        summary = f"共检测到 {len(issues)} 个排查点，主要是模板指导词或合理雷同，风险较低。"
                        
                    collusion_data[competitor_name] = {
                        "summary": summary,
                        "issues": issues
                    }
                global LATEST_COLLUSION_DATA
                LATEST_COLLUSION_DATA = collusion_data
                
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": True, "collusionData": collusion_data}).encode('utf-8'))
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": str(e)}).encode('utf-8'))
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)
            return


    def load_key(self, env_key):
        if not env_key:
            return ""
        
        # 1. Try reading from project .env
        env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
        if os.path.exists(env_path):
            try:
                with open(env_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                match = re.search(rf'^{env_key}=(.*)', content, re.M)
                if match:
                    val = match.group(1).strip().strip("'").strip('"')
                    if val: return val
            except:
                pass

        # 2. Fallback to ~/.baoyu-skills/.env
        backup_env = os.path.expanduser("~/.baoyu-skills/.env")
        if os.path.exists(backup_env):
            try:
                with open(backup_env, 'r', encoding='utf-8') as f:
                    content = f.read()
                match = re.search(rf'^{env_key}=(.*)', content, re.M)
                if match:
                    val = match.group(1).strip().strip("'").strip('"')
                    if val: return val
            except:
                pass

        return os.environ.get(env_key, "")

    def save_key(self, env_key, new_val):
        if not env_key:
            return False, "Invalid env key"
        
        # Save to ~/.baoyu-skills/.env to sync with unified config
        env_path = os.path.expanduser("~/.baoyu-skills/.env")
        try:
            content = ""
            if os.path.exists(env_path):
                with open(env_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            def set_env_val(key, val, current):
                regex = re.compile(rf'^{key}=.*', re.M)
                if regex.search(current):
                    return regex.sub(f'{key}={val}', current)
                prefix = "\n" if current and not current.endswith("\n") else ""
                return current + f'{prefix}{key}={val}'

            new_content = set_env_val(env_key, new_val, content)
            
            os.makedirs(os.path.dirname(env_path), exist_ok=True)
            with open(env_path, 'w', encoding='utf-8') as f:
                f.write(new_content.strip() + '\n')
            
            os.environ[env_key] = new_val
            return True, ""
        except Exception as e:
            return False, str(e)

class ReuseAddrTCPServer(socketserver.TCPServer):
    allow_reuse_address = True

if __name__ == "__main__":
    print("BidShield UI server starting at http://localhost:8000...")
    webbrowser.open("http://localhost:8000")

    # Change working directory to the script's directory so it serves the files correctly
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    with ReuseAddrTCPServer(("", PORT), CustomHandler) as httpd:
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("\nServer stopped.")
