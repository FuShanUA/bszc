import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_formatted_text(paragraph, text):
    # Splits by double asterisks to handle bold text
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.bold = is_bold
        is_bold = not is_bold

def create_docx_table(doc, table_data):
    if not table_data:
        return
        
    rows = len(table_data)
    cols = len(table_data[0])
    
    table = doc.add_table(rows=rows, cols=cols)
    set_table_borders(table)
    
    # Margins are 1" left/right, so content width is 6.5"
    col_widths = [Inches(0.5), Inches(1.0), Inches(0.8), Inches(0.8), Inches(1.0), Inches(2.4)]
    if cols != 6:
        col_widths = [Inches(6.5 / cols)] * cols

    for r_idx, row_data in enumerate(table_data):
        row = table.rows[r_idx]
        
        # Prevent row splitting
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        # Repeat header row on new pages
        if r_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for c_idx, cell_value in enumerate(row_data):
            if c_idx >= cols:
                continue
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            
            # Clean bold markers for table cell text
            clean_value = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_value)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            
            run = p.add_run(clean_value)
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            if r_idx == 0:
                run.font.bold = True
                set_cell_background(cell, "F2F2F2")
            elif r_idx % 2 == 1:
                set_cell_background(cell, "FAFAFA")
                
    # Add spacing paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

def main():
    doc = Document()
    
    # 1" margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Default Normal Style: Arial 11pt
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    md_path = "/Users/shanfu/Downloads/串标分析.md"
    docx_path = "/Users/shanfu/Downloads/串标分析.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: MD file not found: {md_path}")
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            if in_table:
                create_docx_table(doc, table_data)
                table_data = []
                in_table = False
            continue
            
        if stripped.startswith("#"):
            if in_table:
                create_docx_table(doc, table_data)
                table_data = []
                in_table = False
                
            m = re.match(r'^(#+)\s*(.*)', stripped)
            level = len(m.group(1))
            title = m.group(2)
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            
            heading = doc.add_heading(title, level=min(level, 4))
            for run in heading.runs:
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                
        elif stripped.startswith("|"):
            if re.match(r'^\|\s*[:\-]+\s*\|', stripped) or "---" in stripped:
                continue
            in_table = True
            parts = [p.strip() for p in stripped.split("|")[1:-1]]
            table_data.append(parts)
            
        elif stripped.startswith("-") or stripped.startswith("*"):
            if in_table:
                create_docx_table(doc, table_data)
                table_data = []
                in_table = False
                
            text = stripped[1:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_formatted_text(p, text)
            
        else:
            if in_table:
                create_docx_table(doc, table_data)
                table_data = []
                in_table = False
                
            p = doc.add_paragraph()
            parse_formatted_text(p, stripped)
            
    if in_table:
        create_docx_table(doc, table_data)
        
    doc.save(docx_path)
    print(f"Successfully converted MD to DOCX: {docx_path}")

if __name__ == "__main__":
    main()
